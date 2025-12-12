"""
Local File Watcher for processing documents from the new_news folder.
"""

import asyncio
import shutil
from pathlib import Path
from datetime import datetime
from typing import Optional, Callable
from dataclasses import dataclass

import pandas as pd
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler, FileCreatedEvent
from loguru import logger


@dataclass
class ParsedDocument:
    """Parsed document content."""
    file_name: str
    file_type: str
    content: str
    page_count: int = 1
    chunk_count: int = 1
    metadata: dict = None
    structured_data: list = None  # For CSV files with trading data
    
    @property
    def is_structured(self) -> bool:
        return self.structured_data is not None and len(self.structured_data) > 0


def safe_get_symbol(row: pd.Series) -> Optional[str]:
    """
    Safely get symbol from row, handling NaN values.
    
    Args:
        row: DataFrame row
        
    Returns:
        Symbol string or None if not found/NaN
    """
    # Try 'symbol' column first
    symbol = row.get('symbol')
    if symbol is not None and pd.notna(symbol):
        return str(symbol).strip()
    
    # Try 'ticker' column as fallback
    ticker = row.get('ticker')
    if ticker is not None and pd.notna(ticker):
        return str(ticker).strip()
    
    return None


class DocumentParser:
    """Parser for different document types."""
    
    async def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a document based on its extension."""
        ext = file_path.suffix.lower()
        
        if ext == ".pdf":
            return await self._parse_pdf(file_path)
        elif ext == ".docx":
            return await self._parse_docx(file_path)
        elif ext in (".csv", ".xlsx"):
            return await self._parse_spreadsheet(file_path)
        elif ext == ".txt":
            return await self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    async def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse PDF document."""
        try:
            from pypdf import PdfReader  # type: ignore[import-untyped,import-not-found]
            
            loop = asyncio.get_event_loop()
            
            def extract():
                reader = PdfReader(file_path)
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                return pages, len(reader.pages)
            
            pages, page_count = await loop.run_in_executor(None, extract)
            content = "\n\n".join(pages)
            
            return ParsedDocument(
                file_name=file_path.name,
                file_type="pdf",
                content=content,
                page_count=page_count,
                metadata={"source": "local_file"},
            )
            
        except ImportError:
            logger.error("pypdf not installed, cannot parse PDF")
            raise
    
    async def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse Word document."""
        try:
            from docx import Document  # type: ignore[import-untyped]
            
            loop = asyncio.get_event_loop()
            
            def extract():
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return paragraphs
            
            paragraphs = await loop.run_in_executor(None, extract)
            content = "\n\n".join(paragraphs)
            
            return ParsedDocument(
                file_name=file_path.name,
                file_type="docx",
                content=content,
                page_count=1,
                metadata={"source": "local_file"},
            )
            
        except ImportError:
            logger.error("python-docx not installed, cannot parse DOCX")
            raise
    
    async def _parse_spreadsheet(self, file_path: Path) -> ParsedDocument:
        """Parse CSV or Excel spreadsheet."""
        try:
            loop = asyncio.get_event_loop()
            
            def extract():
                if file_path.suffix.lower() == ".csv":
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                return df
            
            df = await loop.run_in_executor(None, extract)
            
            # Check if it's structured trading data
            trading_columns = {'symbol', 'ticker', 'rating', 'price_target', 'action', 'sentiment'}
            has_trading_data = bool(trading_columns & set(df.columns.str.lower()))
            
            if has_trading_data:
                # Normalize column names
                df.columns = df.columns.str.lower()
                
                # Extract structured data
                structured_data = []
                for _, row in df.iterrows():
                    # Use safe_get_symbol to handle NaN values properly
                    symbol = safe_get_symbol(row)
                    if symbol:
                        # Safely extract other fields, handling NaN
                        structured_data.append({
                            "symbol": symbol.upper(),
                            "rating": row.get('rating') if pd.notna(row.get('rating')) else None,
                            "price_target": float(row.get('price_target')) if pd.notna(row.get('price_target')) else None,
                            "action": str(row.get('action')) if pd.notna(row.get('action')) else None,
                            "sentiment": float(row.get('sentiment')) if pd.notna(row.get('sentiment')) else None,
                            "analyst": str(row.get('analyst')) if pd.notna(row.get('analyst')) else None,
                            "source": str(row.get('source')) if pd.notna(row.get('source')) else None,
                        })
                
                return ParsedDocument(
                    file_name=file_path.name,
                    file_type="csv_structured",
                    content=df.to_string(),
                    structured_data=structured_data,
                    metadata={"columns": list(df.columns)},
                )
            
            return ParsedDocument(
                file_name=file_path.name,
                file_type="csv",
                content=df.to_string(),
                metadata={"columns": list(df.columns), "rows": len(df)},
            )
            
        except ImportError:
            logger.error("pandas not installed, cannot parse spreadsheet")
            raise
    
    async def _parse_text(self, file_path: Path) -> ParsedDocument:
        """Parse plain text file."""
        loop = asyncio.get_event_loop()
        
        def read():
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        content = await loop.run_in_executor(None, read)
        
        return ParsedDocument(
            file_name=file_path.name,
            file_type="txt",
            content=content,
            metadata={"source": "local_file"},
        )


class LocalFileHandler(FileSystemEventHandler):
    """Watchdog event handler for new files."""
    
    def __init__(self, queue: asyncio.Queue, loop: asyncio.AbstractEventLoop):
        self.queue = queue
        self.loop = loop
        self.supported_extensions = {'.csv', '.xlsx', '.pdf', '.docx', '.txt'}
    
    def on_created(self, event):
        if event.is_directory:
            return
        
        file_path = Path(event.src_path)
        if file_path.suffix.lower() in self.supported_extensions:
            # Schedule async task
            asyncio.run_coroutine_threadsafe(
                self.queue.put(file_path),
                self.loop
            )
            logger.info(f"New file detected: {file_path.name}")


class LocalFileWatcher:
    """
    Watch the new_news folder for new documents and process them.
    
    Folder structure:
    - new_news/incoming/  - Drop files here
    - new_news/processed/ - Successfully processed files
    - new_news/failed/    - Files that failed processing
    """
    
    def __init__(
        self,
        base_path: str = "new_news",
        on_document: Callable = None,
    ):
        """
        Initialize file watcher.
        
        Args:
            base_path: Base path for news folders
            on_document: Callback for processed documents
        """
        self.base_path = Path(base_path)
        self.incoming_path = self.base_path / "incoming"
        self.processed_path = self.base_path / "processed"
        self.failed_path = self.base_path / "failed"
        
        self.on_document = on_document
        self.parser = DocumentParser()
        self.queue: asyncio.Queue = None
        self.observer: Observer = None
        self._running = False
    
    async def start(self) -> None:
        """Start watching for files."""
        # Create directories
        for path in [self.incoming_path, self.processed_path, self.failed_path]:
            path.mkdir(parents=True, exist_ok=True)
        
        # Create queue
        self.queue = asyncio.Queue()
        loop = asyncio.get_event_loop()
        
        # Setup watchdog
        handler = LocalFileHandler(self.queue, loop)
        self.observer = Observer()
        self.observer.schedule(handler, str(self.incoming_path), recursive=False)
        self.observer.start()
        
        self._running = True
        logger.info(f"File watcher started, monitoring: {self.incoming_path}")
        
        # Process existing files
        await self._process_existing_files()
        
        # Start processing loop
        asyncio.create_task(self._process_loop())
    
    async def stop(self) -> None:
        """Stop watching for files."""
        self._running = False
        if self.observer:
            self.observer.stop()
            self.observer.join()
        logger.info("File watcher stopped")
    
    async def _process_existing_files(self) -> None:
        """Process any existing files in incoming folder."""
        for file_path in self.incoming_path.iterdir():
            if file_path.is_file():
                await self.queue.put(file_path)
    
    async def _process_loop(self) -> None:
        """Main processing loop."""
        while self._running:
            try:
                # Wait for file with timeout
                try:
                    file_path = await asyncio.wait_for(
                        self.queue.get(),
                        timeout=1.0
                    )
                except asyncio.TimeoutError:
                    continue
                
                # Process the file
                await self._process_file(file_path)
                
            except Exception as e:
                logger.error(f"Error in processing loop: {e}")
    
    async def _process_file(self, file_path: Path) -> None:
        """Process a single file."""
        start_time = datetime.now()
        
        try:
            logger.info(f"Processing file: {file_path.name}")
            
            # Parse the document
            document = await self.parser.parse(file_path)
            
            # Call callback if provided
            if self.on_document:
                await self.on_document(document)
            
            # Move to processed
            dest = self.processed_path / file_path.name
            shutil.move(str(file_path), str(dest))
            
            processing_time = (datetime.now() - start_time).total_seconds() * 1000
            logger.info(
                f"Processed {file_path.name} in {processing_time:.0f}ms "
                f"({document.page_count} pages)"
            )
            
        except Exception as e:
            logger.error(f"Failed to process {file_path.name}: {e}")
            
            # Move to failed
            try:
                dest = self.failed_path / file_path.name
                shutil.move(str(file_path), str(dest))
            except Exception:
                pass
    
    def chunk_content(self, content: str, max_tokens: int = 4000) -> list[str]:
        """Split content into chunks for processing."""
        # Rough estimate: 4 characters per token
        max_chars = max_tokens * 4
        
        if len(content) <= max_chars:
            return [content]
        
        chunks = []
        paragraphs = content.split('\n\n')
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) <= max_chars:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks

